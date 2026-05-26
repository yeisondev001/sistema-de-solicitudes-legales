from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .base import (
    doc_base, tabla_encabezado, fila, run, p, parse_fecha, fecha_larga, a_bytes,
)


def generar(d: dict) -> bytes:
    doc = doc_base()
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    C = WD_ALIGN_PARAGRAPH.CENTER

    fecha = parse_fecha(d["fecha_carta"]) or date.today()
    p(doc, "Santo Domingo, D. N.", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(doc, fecha_larga(fecha) + ".", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(doc)

    t = tabla_encabezado(doc)
    fila(t, "A",
         (d["destinatario"],       True),
         (d["cargo_destinatario"], False))
    fila(t, "Atención",
         (d["atencion_nombre"], True),
         (d["atencion_cargo"],  False))
    fila(t, "Asunto", (d["asunto"], True))
    p(doc)

    ps = doc.add_paragraph()
    ps.alignment = J
    ps.paragraph_format.space_after = ps.paragraph_format.space_before = Pt(0)
    run(ps, "Distinguida señora:", bold=True)
    p(doc)

    p1 = doc.add_paragraph()
    p1.alignment = J
    p1.paragraph_format.space_after = p1.paragraph_format.space_before = Pt(0)
    run(p1, d["cuerpo_1"])

    if d.get("cuerpo_2", "").strip():
        p(doc)
        p2 = doc.add_paragraph()
        p2.alignment = J
        p2.paragraph_format.space_after = p2.paragraph_format.space_before = Pt(0)
        run(p2, d["cuerpo_2"])

    p(doc)
    p(doc, "Con saludos de alta estima;", align=J)
    p(doc)

    pc = doc.add_paragraph()
    pc.alignment = C
    pc.paragraph_format.space_after = pc.paragraph_format.space_before = Pt(0)
    run(pc, "Deferentemente;")

    # Firma manual (sin bloque estándar para incluir Cc)
    p(doc); p(doc)
    for linea in [
        ("______________________________________________", False),
        (d["firmante_nombre"], True),
        (d["firmante_cargo"],  False),
    ]:
        par = doc.add_paragraph()
        par.alignment = C
        par.paragraph_format.space_after = par.paragraph_format.space_before = Pt(0)
        pPr = par._p.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:line"), "276"); sp.set(qn("w:lineRule"), "auto")
        pPr.append(sp)
        run(par, linea[0], bold=linea[1])

    p(doc)
    p(doc, d["iniciales"])
    if d.get("cc", "").strip():
        p(doc, f"Cc. {d['cc']}")

    return a_bytes(doc)
