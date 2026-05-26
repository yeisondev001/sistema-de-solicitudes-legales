"""
Constantes y helpers compartidos por todos los generadores de documentos Word.
"""

import io, re
from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FUENTE_DOC = "Century Gothic"
COL_WIDTHS = (2579, 449, 5692)

MESES = {
    1:"Enero", 2:"Febrero", 3:"Marzo",    4:"Abril",
    5:"Mayo",  6:"Junio",   7:"Julio",    8:"Agosto",
    9:"Septiembre", 10:"Octubre", 11:"Noviembre", 12:"Diciembre",
}


def fecha_larga(d: date) -> str:
    return f"{d.day} de {MESES[d.month]} de {d.year}"


def parse_fecha(txt: str):
    try:
        p = txt.strip().split("/")
        if len(p) == 3:
            return date(int(p[2]), int(p[1]), int(p[0]))
    except Exception:
        pass
    return None


def run(p, texto, bold=False, size=12):
    r = p.add_run(texto)
    r.font.name = FUENTE_DOC
    r.font.size = Pt(size)
    r.bold = bold
    return r


def p(doc, texto="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.space_after  = Pt(0)
    par.paragraph_format.space_before = Pt(0)
    if texto:
        run(par, texto, bold=bold, size=size)
    return par


def set_cell_w(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    for o in tcPr.findall(qn("w:tcW")):
        tcPr.remove(o)
    e = OxmlElement("w:tcW")
    e.set(qn("w:w"), str(dxa)); e.set(qn("w:type"), "dxa")
    tcPr.append(e)


def tabla_encabezado(doc):
    tabla = doc.add_table(rows=0, cols=3)
    tbl   = tabla._tbl
    tblPr = tbl.tblPr
    for o in tblPr.findall(qn("w:tblW")):
        tblPr.remove(o)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(sum(COL_WIDTHS))); tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in [("top","0"),("left","108"),("bottom","0"),("right","108")]:
        mar = OxmlElement(f"w:{side}")
        mar.set(qn("w:w"), val); mar.set(qn("w:type"), "dxa")
        tblCellMar.append(mar)
    tblPr.append(tblCellMar)
    tb = OxmlElement("w:tblBorders")
    for lado in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{lado}")
        b.set(qn("w:val"),"none"); b.set(qn("w:sz"),"0")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"auto")
        tb.append(b)
    tblPr.append(tb)
    for o in tbl.findall(qn("w:tblGrid")):
        tbl.remove(o)
    tg = OxmlElement("w:tblGrid")
    for w in COL_WIDTHS:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); tg.append(gc)
    tblPr.addnext(tg)
    return tabla


def fila(tabla, etiqueta, *valores):
    fila_ = tabla.add_row()
    c1, c2, c3 = fila_.cells
    for cell, dxa in zip([c1,c2,c3], COL_WIDTHS):
        set_cell_w(cell, dxa)
    c1.paragraphs[0].clear()
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0); p1.paragraph_format.space_before = Pt(0)
    run(p1, etiqueta, bold=True)
    c2.paragraphs[0].clear()
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.space_before = Pt(0)
    run(p2, ":", bold=True)
    c3.paragraphs[0].clear()
    primera, *resto = valores
    px = c3.paragraphs[0]
    px.paragraph_format.space_after = Pt(0); px.paragraph_format.space_before = Pt(0)
    run(px, primera[0], bold=primera[1])
    for txt, bld in resto:
        py = c3.add_paragraph()
        py.paragraph_format.space_after = Pt(0); py.paragraph_format.space_before = Pt(0)
        run(py, txt, bold=bld)
    c3.paragraphs[-1].paragraph_format.space_after = Pt(8)
    p1.paragraph_format.space_after = Pt(8)


def footer_iniciales(doc, iniciales):
    section = doc.sections[0]
    section.footer_distance = Cm(1.5)
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    run(fp, iniciales, size=10)


def doc_base():
    document = Document()
    for s in document.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3)
        s.right_margin  = Cm(2.5)
    document.styles["Normal"].font.name = FUENTE_DOC
    document.styles["Normal"].font.size = Pt(12)
    return document


def firma_bloque(doc, d):
    def pc(texto, bold=False):
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after  = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        pPr = par._p.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:line"), "276"); sp.set(qn("w:lineRule"), "auto")
        pPr.append(sp)
        run(par, texto, bold=bold)
    p(doc); p(doc)
    pc("______________________________________________")
    pc(d["firmante_nombre"], bold=True)
    pc(d["firmante_cargo"])


def nombre_archivo(tipo: str, identificador: str) -> str:
    limpio = re.sub(r'[\\/*?:"<>|]', '', identificador).strip()
    limpio = re.sub(r'\s+', ' ', limpio)
    return f"{tipo} - {limpio}.docx"


def a_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
