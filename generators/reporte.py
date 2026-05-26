from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .base import (
    doc_base, tabla_encabezado, fila, run, p, parse_fecha, MESES, a_bytes,
)


def generar(d: dict) -> bytes:
    doc = doc_base()
    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    L = WD_ALIGN_PARAGRAPH.LEFT

    # Título
    t = doc.add_paragraph()
    t.alignment = C
    t.paragraph_format.space_after = t.paragraph_format.space_before = Pt(0)
    r = t.add_run("REPORTE DE AUDIENCIA")
    r.font.name = "Century Gothic"; r.font.size = Pt(13)
    r.bold = True; r.underline = True
    p(doc)

    # Caja / Rol / Fecha
    fa = parse_fecha(d["fecha_audiencia"])
    fecha_txt = f"{fa.day} de {MESES[fa.month]} {fa.year}" if fa else d["fecha_audiencia"]
    info = doc.add_paragraph()
    info.alignment = J
    info.paragraph_format.space_after = info.paragraph_format.space_before = Pt(0)
    run(info, "Caja  No.", bold=True); run(info, f"  {d['caja_no']}          ")
    run(info, "Rol",       bold=True); run(info, f"  {d['rol']}           ")
    run(info, "Fecha ",    bold=True); run(info, fecha_txt)
    p(doc)

    # Tabla encabezado
    tabla = tabla_encabezado(doc)
    fila(tabla, "Tribunal",       (d["tribunal"],    False))
    fila(tabla, "No. Expediente", (d["expediente"],  False))
    fila(tabla, "Demandante(s)",  (d["demandantes"], False))
    demandados = [x.strip() for x in d.get("demandados", []) if x.strip()]
    if demandados:
        fila(tabla, "Demandado(s)", (demandados[0], False))
        for dem in demandados[1:]:
            fila(tabla, "Demandado", (dem, False))
    fila(tabla, "Asunto",     (d["asunto"],  False))
    fila(tabla, "Parcela No", (d["parcela"], False))
    p(doc)

    # Demandante Conclusiones
    pd = doc.add_paragraph()
    pd.alignment = C
    pd.paragraph_format.space_after = pd.paragraph_format.space_before = Pt(0)
    run(pd, "Demandante", bold=True)
    run(pd, " Conclusiones:" + "_" * 168)
    p(doc)

    # Demandado + Conclusiones
    pd2 = doc.add_paragraph()
    pd2.alignment = C
    pd2.paragraph_format.space_after = pd2.paragraph_format.space_before = Pt(0)
    run(pd2, "Demandado", bold=True)
    pc = doc.add_paragraph()
    pc.alignment = L
    pc.paragraph_format.space_after = pc.paragraph_format.space_before = Pt(0)
    run(pc, "Conclusiones:" + "_" * 168)
    p(doc)

    # Fallo
    pf = doc.add_paragraph()
    pf.alignment = C
    pf.paragraph_format.space_after = pf.paragraph_format.space_before = Pt(0)
    run(pf, "Fallo:", bold=True)
    p(doc)
    pft = doc.add_paragraph()
    pft.alignment = L
    pft.paragraph_format.space_after = pft.paragraph_format.space_before = Pt(0)
    run(pft, d["fallo"])
    p(doc)

    # Próxima audiencia
    fp_obj = parse_fecha(d["fecha_proxima"])
    fp_txt = f"{fp_obj.day} de {MESES[fp_obj.month]} de {fp_obj.year}" if fp_obj else d["fecha_proxima"]
    pfp = doc.add_paragraph()
    pfp.alignment = L
    pfp.paragraph_format.space_after = pfp.paragraph_format.space_before = Pt(0)
    run(pfp, "Fecha de la próxima audiencia:", bold=True)
    run(pfp, f" {fp_txt}.")
    p(doc)

    # Abogado
    pab = doc.add_paragraph()
    pab.alignment = L
    pab.paragraph_format.space_after = pab.paragraph_format.space_before = Pt(0)
    run(pab, "Abogado que Compareció:", bold=True)
    run(pab, f" {d['abogados']}.")

    return a_bytes(doc)
