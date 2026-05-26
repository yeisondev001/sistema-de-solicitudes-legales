"""
Automatización de Solicitudes Legales
======================================
Dos plantillas seleccionables:
  1. Solicitud de Informe Técnico (Catastro)
  2. Solicitud de Copia Certificada de Expediente (Archivo)
Frontend: navy + dorado · Cormorant Garamond / Source Serif 4
"""

import flet as ft
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import io, os, platform, base64, pathlib

# ──────────────────────────────────────────────
# PALETA  Navy & Oro
# ──────────────────────────────────────────────
INK        = "#0F2A3F"
INK_SOFT   = "#2C4A63"
ACCENT     = "#B8860B"
PAPER      = "#FBF7EE"
PAPER_EDGE = "#EFE7D2"
RULE       = "#D9C998"
SURFACE    = "#F7F2E6"
PANEL_INK  = "#0F2A3F"

FONT_DISPLAY = "Cormorant Garamond"
FONT_BODY    = "Source Serif 4"

# ──────────────────────────────────────────────
# CONSTANTES DEL DOCUMENTO
# ──────────────────────────────────────────────
FUENTE_DOC = "Century Gothic"
COL_WIDTHS = (2579, 449, 5692)

MESES = {
    1:"Enero",2:"Febrero",3:"Marzo",4:"Abril",5:"Mayo",6:"Junio",
    7:"Julio",8:"Agosto",9:"Septiembre",10:"Octubre",11:"Noviembre",12:"Diciembre",
}

# ══════════════════════════════════════════════
#  HELPERS WORD
# ══════════════════════════════════════════════

def _fecha_larga(d: date) -> str:
    return f"{d.day} de {MESES[d.month]} de {d.year}"

def _parse_fecha(txt: str):
    try:
        p = txt.strip().split("/")
        if len(p) == 3:
            return date(int(p[2]), int(p[1]), int(p[0]))
    except Exception:
        pass
    return None

def _run(p, texto, bold=False, size=12):
    r = p.add_run(texto)
    r.font.name = FUENTE_DOC
    r.font.size = Pt(size)
    r.bold = bold
    return r

def _p(doc, texto="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.space_after  = Pt(0)
    par.paragraph_format.space_before = Pt(0)
    if texto:
        _run(par, texto, bold=bold, size=size)
    return par

def _set_cell_w(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    for o in tcPr.findall(qn("w:tcW")):
        tcPr.remove(o)
    e = OxmlElement("w:tcW")
    e.set(qn("w:w"), str(dxa)); e.set(qn("w:type"), "dxa")
    tcPr.append(e)

def _tabla_encabezado(doc):
    tabla = doc.add_table(rows=0, cols=3)
    tbl   = tabla._tbl
    tblPr = tbl.tblPr
    for o in tblPr.findall(qn("w:tblW")):
        tblPr.remove(o)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(sum(COL_WIDTHS))); tw.set(qn("w:type"), "dxa")
    tblPr.append(tw)
    # Márgenes celda: sin espacio arriba/abajo
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in [("top","0"),("left","108"),("bottom","0"),("right","108")]:
        mar = OxmlElement(f"w:{side}")
        mar.set(qn("w:w"), val); mar.set(qn("w:type"), "dxa")
        tblCellMar.append(mar)
    tblPr.append(tblCellMar)
    # Sin bordes
    tb = OxmlElement("w:tblBorders")
    for lado in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{lado}")
        b.set(qn("w:val"),"none"); b.set(qn("w:sz"),"0")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"auto")
        tb.append(b)
    tblPr.append(tb)
    # Grilla
    for o in tbl.findall(qn("w:tblGrid")):
        tbl.remove(o)
    tg = OxmlElement("w:tblGrid")
    for w in COL_WIDTHS:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); tg.append(gc)
    tblPr.addnext(tg)
    return tabla

def _fila(tabla, etiqueta, *valores):
    fila = tabla.add_row()
    c1, c2, c3 = fila.cells
    for cell, dxa in zip([c1,c2,c3], COL_WIDTHS):
        _set_cell_w(cell, dxa)
    c1.paragraphs[0].clear()
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0); p1.paragraph_format.space_before = Pt(0)
    _run(p1, etiqueta, bold=True)
    c2.paragraphs[0].clear()
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.space_before = Pt(0)
    _run(p2, ":", bold=True)
    c3.paragraphs[0].clear()
    primera, *resto = valores
    px = c3.paragraphs[0]
    px.paragraph_format.space_after = Pt(0); px.paragraph_format.space_before = Pt(0)
    _run(px, primera[0], bold=primera[1])
    for txt, bld in resto:
        py = c3.add_paragraph()
        py.paragraph_format.space_after = Pt(0); py.paragraph_format.space_before = Pt(0)
        _run(py, txt, bold=bld)
    c3.paragraphs[-1].paragraph_format.space_after = Pt(8)
    p1.paragraph_format.space_after = Pt(8)

def _footer_iniciales(doc, iniciales):
    section = doc.sections[0]
    section.footer_distance = Cm(1.5)
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    _run(fp, iniciales, size=10)

def _doc_base():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3)
        s.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = FUENTE_DOC
    doc.styles["Normal"].font.size = Pt(12)
    return doc

def _firma_bloque(doc, d):
    def pc(texto, bold=False):
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after  = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        pPr = par._p.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:line"), "276"); sp.set(qn("w:lineRule"), "auto")
        pPr.append(sp)
        _run(par, texto, bold=bold)
    _p(doc); _p(doc)
    pc("______________________________________________")
    pc(d["firmante_nombre"], bold=True)
    pc(d["firmante_cargo"])

# ──────────────────────────────────────────────
# PLANTILLA 1: INFORME TÉCNICO (Catastro)
# ──────────────────────────────────────────────

def _cuerpo_informe(doc, d):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after  = Pt(0)
    _run(par, "\t\t")
    _run(par, "Cortésmente, tenemos a bien solicitarle interponer de sus buenos oficios, "
              "a los fines de que el Departamento de Catastro de esta institución, realice una "
              "Investigación Parcelaria, sobre el inmueble identificado como: «")
    _run(par, f" Inscripción Catastral No. {d['inscripcion_no']} , Expedido por la Dirección "
              f"General de Catastro Nacional a nombre de {d['propietario']} »", bold=True)
    _run(par, ", donde se establezca si el ")
    _run(par, "Estado Dominicano", bold=True)
    _run(par, " tiene o no derechos sobre el inmueble anteriormente descrito "
              "o si existe alguna anotación a favor de este, a los fines de ser utilizado "
              "como medio probatorio ")
    _run(par, f"para la audiencia fijada para el {d['fecha_audiencia']},", bold=True)
    _run(par, f" por ante el Salón {d['salon']}, {d['piso']}, {d['lugar']}")
    _run(par, f", Exp. {d['expediente']}", bold=True)
    _run(par, ", a nombre de la señora ")
    _run(par, f"{d['cliente']}.", bold=True)

def generar_informe_tecnico(d: dict) -> bytes:
    doc = _doc_base()
    fecha = _parse_fecha(d["fecha_carta"]) or date.today()
    _p(doc, "Santo Domingo, D.N.", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, _fecha_larga(fecha),   bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc)
    tabla = _tabla_encabezado(doc)
    _fila(tabla, "Al",
          (d["destinatario"],       True),
          (d["cargo_destinatario"], False))
    _fila(tabla, "Atención",
          (d["atencion_nombre"], True),
          (d["atencion_cargo"],  False))
    _fila(tabla, "Asunto",
          (d["asunto"], False))
    anexos = [a.strip() for a in d.get("anexos", []) if a.strip()]
    if anexos:
        _fila(tabla, "Anexo",
              *[(f"{i+1}- {a}", False) for i, a in enumerate(anexos)])
    _p(doc)
    _cuerpo_informe(doc, d)
    _p(doc)
    _p(doc, "Nota: Esta solicitud de Informe Técnico, está exento de pago.")
    _firma_bloque(doc, d)
    _footer_iniciales(doc, d["iniciales"])
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ──────────────────────────────────────────────
# PLANTILLA 2: COPIA CERTIFICADA (Archivo)
# ──────────────────────────────────────────────

def _cuerpo_archivo(doc, d):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after  = Pt(0)
    _run(par, "\t\t")
    _run(par, "Cortésmente, tenemos a bien solicitar que interponga sus buenos oficios "
              "con el objetivo de remitirnos una Copia Certificada del Expediente "
              "completo a nombre del señor ")
    _run(par, f"{d['nombre_cliente']}", bold=True)
    _run(par, ", teniendo de la cédula de identidad y electoral No. ")
    _run(par, f"{d['cedula']}")
    _run(par, ", en relación con el inmueble identificado como: «")
    _run(par, f"{d['descripcion_inmueble']}", bold=True)
    _run(par, f"», a los fines de ser utilizado como medio probatorio en el proceso de ")
    _run(par, f"{d['tipo_proceso']}", bold=True)
    _run(par, " que se está conociendo por ante la ")
    _run(par, f"{d['tribunal']}", bold=True)
    _run(par, " ».")

def generar_copia_certificada(d: dict) -> bytes:
    doc = _doc_base()
    fecha = _parse_fecha(d["fecha_carta"]) or date.today()
    _p(doc, "Santo Domingo, D. N.", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, _fecha_larga(fecha) + ".",  bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc)
    tabla = _tabla_encabezado(doc)
    _fila(tabla, "A",
          (d["destinatario"],       True),
          (d["cargo_destinatario"], False))
    _fila(tabla, "Asunto",
          (d["asunto"], False))
    _p(doc)
    _cuerpo_archivo(doc, d)
    _p(doc)
    _p(doc, "Atentamente,", align=WD_ALIGN_PARAGRAPH.CENTER)
    _firma_bloque(doc, d)
    _footer_iniciales(doc, d["iniciales"])
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ──────────────────────────────────────────────
# PLANTILLA 3: REPORTE DE AUDIENCIA
# ──────────────────────────────────────────────

def generar_reporte_audiencia(d: dict) -> bytes:
    doc = _doc_base()
    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY
    L = WD_ALIGN_PARAGRAPH.LEFT

    # Título subrayado
    t = doc.add_paragraph()
    t.alignment = C
    t.paragraph_format.space_after  = Pt(0)
    t.paragraph_format.space_before = Pt(0)
    r = t.add_run("REPORTE DE AUDIENCIA")
    r.font.name = FUENTE_DOC
    r.font.size = Pt(13)
    r.bold = True
    r.underline = True
    _p(doc)

    # Línea Caja / Rol / Fecha (ANTES de la tabla)
    fa = _parse_fecha(d["fecha_audiencia"])
    fecha_txt = f"{fa.day} de {MESES[fa.month]} {fa.year}" if fa else d["fecha_audiencia"]
    info = doc.add_paragraph()
    info.alignment = J
    info.paragraph_format.space_after  = Pt(0)
    info.paragraph_format.space_before = Pt(0)
    _run(info, "Caja  No.", bold=True)
    _run(info, f"  {d['caja_no']}          ")
    _run(info, "Rol", bold=True)
    _run(info, f"  {d['rol']}           ")
    _run(info, "Fecha ", bold=True)
    _run(info, fecha_txt)
    _p(doc)

    # Tabla encabezado
    tabla = _tabla_encabezado(doc)
    _fila(tabla, "Tribunal",       (d["tribunal"],    False))
    _fila(tabla, "No. Expediente", (d["expediente"],  False))
    _fila(tabla, "Demandante(s)",  (d["demandantes"], False))
    demandados = [x.strip() for x in d.get("demandados", []) if x.strip()]
    if demandados:
        _fila(tabla, "Demandado(s)", (demandados[0], False))
        for dem in demandados[1:]:
            _fila(tabla, "Demandado", (dem, False))
    _fila(tabla, "Asunto",     (d["asunto"],  False))
    _fila(tabla, "Parcela No", (d["parcela"], False))
    _p(doc)

    # Demandante Conclusiones (misma línea)
    pd = doc.add_paragraph()
    pd.alignment = C
    pd.paragraph_format.space_after  = Pt(0)
    pd.paragraph_format.space_before = Pt(0)
    _run(pd, "Demandante", bold=True)
    _run(pd, " Conclusiones:" + "_" * 168)
    _p(doc)

    # Demandado + Conclusiones
    pd2 = doc.add_paragraph()
    pd2.alignment = C
    pd2.paragraph_format.space_after  = Pt(0)
    pd2.paragraph_format.space_before = Pt(0)
    _run(pd2, "Demandado", bold=True)

    pc = doc.add_paragraph()
    pc.alignment = L
    pc.paragraph_format.space_after  = Pt(0)
    pc.paragraph_format.space_before = Pt(0)
    _run(pc, "Conclusiones:" + "_" * 168)
    _p(doc)

    # Fallo
    pf = doc.add_paragraph()
    pf.alignment = C
    pf.paragraph_format.space_after  = Pt(0)
    pf.paragraph_format.space_before = Pt(0)
    _run(pf, "Fallo:", bold=True)
    _p(doc)

    # Texto del fallo — alineado izquierda
    pft = doc.add_paragraph()
    pft.alignment = L
    pft.paragraph_format.space_after  = Pt(0)
    pft.paragraph_format.space_before = Pt(0)
    _run(pft, d["fallo"])
    _p(doc)

    # Fecha próxima audiencia
    fp_obj = _parse_fecha(d["fecha_proxima"])
    fp_txt = _fecha_larga(fp_obj) if fp_obj else d["fecha_proxima"]
    pfp = doc.add_paragraph()
    pfp.alignment = L
    pfp.paragraph_format.space_after  = Pt(0)
    pfp.paragraph_format.space_before = Pt(0)
    _run(pfp, "Fecha de la próxima audiencia:", bold=True)
    _run(pfp, f" {fp_txt}.")
    _p(doc)

    # Abogado
    pab = doc.add_paragraph()
    pab.alignment = L
    pab.paragraph_format.space_after  = Pt(0)
    pab.paragraph_format.space_before = Pt(0)
    _run(pab, "Abogado que Compareció:", bold=True)
    _run(pab, f" {d['abogados']}.")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ──────────────────────────────────────────────
# PLANTILLA 4: SOLICITUD DE COBROS
# ──────────────────────────────────────────────

def generar_solicitud_cobros(d: dict) -> bytes:
    doc = _doc_base()
    # Márgenes iguales a original (right=3cm)
    for s in doc.sections:
        s.right_margin = Cm(3)

    fecha = _parse_fecha(d["fecha_carta"]) or date.today()
    _p(doc, "Santo Domingo DN,", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, _fecha_larga(fecha) + ".", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc)

    # Encabezado con tabla
    tabla = _tabla_encabezado(doc)
    _fila(tabla, "A LA",
          (d["destinatario"],       True),
          (d["cargo_destinatario"], False))
    _fila(tabla, "ASUNTO",
          (d["asunto"], False))
    _p(doc)

    _p(doc, "Distinguida señora:", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _p(doc)

    # Cuerpo
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after  = Pt(0)
    _run(par, "\t\t")
    _run(par, "Cortésmente, tenemos a bien solicitarle, a la mayor brevedad posible, "
              "si existe algún pago de compra a nombre de los Señores o entidad Jurídica "
              "con relación al señor ")
    _run(par, d["nombre_cliente"], bold=True)
    _run(par, f", para conocer del {d['tipo_proceso']}, dentro de la "
              f"{d['parcela']}, del Municipio {d['municipio']} de la provincia "
              f"{d['provincia']}, EXP. N0.{d['expediente']}, Ced. No. {d['cedula']}, "
              f"en el {d['tribunal']}, contra la sentencia N0.{d['sentencia']}, "
              f"de fecha {d['fecha_sentencia']}, dictada por el {d['tribunal_original']}.")
    _p(doc)
    _p(doc, "Con saludos de alta estima;")
    _firma_bloque(doc, d)
    _footer_iniciales(doc, d["iniciales"])
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ──────────────────────────────────────────────
# PLANTILLA 5: SOLICITUD DE INSPECCIÓN
# ──────────────────────────────────────────────

def generar_solicitud_inspeccion(d: dict) -> bytes:
    doc = _doc_base()
    C = WD_ALIGN_PARAGRAPH.CENTER
    J = WD_ALIGN_PARAGRAPH.JUSTIFY

    fecha = _parse_fecha(d["fecha_carta"]) or date.today()
    _p(doc, "Santo Domingo, D. N.", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, _fecha_larga(fecha) + ".", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc)

    # Encabezado con tabla (valores en negrita)
    tabla = _tabla_encabezado(doc)
    _fila(tabla, "A",
          (d["destinatario"],       True),
          (d["cargo_destinatario"], False))
    _fila(tabla, "Atención",
          (d["atencion_nombre"], True),
          (d["atencion_cargo"],  False))
    _fila(tabla, "Asunto",
          (d["asunto"], True))
    _p(doc)

    # Saludo
    ps = doc.add_paragraph()
    ps.alignment = J
    ps.paragraph_format.space_after  = Pt(0)
    ps.paragraph_format.space_before = Pt(0)
    _run(ps, "Distinguida señora:", bold=True)
    _p(doc)

    # Párrafo 1
    p1 = doc.add_paragraph()
    p1.alignment = J
    p1.paragraph_format.space_after  = Pt(0)
    p1.paragraph_format.space_before = Pt(0)
    _run(p1, d["cuerpo_1"])

    # Párrafo 2 (opcional)
    if d.get("cuerpo_2", "").strip():
        _p(doc)
        p2 = doc.add_paragraph()
        p2.alignment = J
        p2.paragraph_format.space_after  = Pt(0)
        p2.paragraph_format.space_before = Pt(0)
        _run(p2, d["cuerpo_2"])

    _p(doc)
    _p(doc, "Con saludos de alta estima;", align=J)

    # Firma con "Deferentemente;"
    _p(doc)
    pc = doc.add_paragraph()
    pc.alignment = C
    pc.paragraph_format.space_after  = Pt(0)
    pc.paragraph_format.space_before = Pt(0)
    _run(pc, "Deferentemente;")
    _firma_bloque(doc, d)

    # Iniciales y Cc en el cuerpo (no footer)
    _p(doc)
    _p(doc, d["iniciales"])
    if d.get("cc", "").strip():
        _p(doc, f"Cc. {d['cc']}")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ══════════════════════════════════════════════
#  HELPER DESCARGA (Windows local / Web Linux)
# ══════════════════════════════════════════════

def _nombre_archivo(tipo: str, identificador: str) -> str:
    """Genera un nombre de archivo limpio: 'Tipo - Nombre Completo.docx'"""
    import re
    limpio = re.sub(r'[\\/*?:"<>|]', '', identificador).strip()
    limpio = re.sub(r'\s+', ' ', limpio)
    return f"{tipo} - {limpio}.docx"

def _guardar_y_abrir(page, contenido: bytes, nombre: str, snack_fn):
    if platform.system() == "Windows":
        # Local: guarda en Descargas y abre con Word
        descargas = pathlib.Path.home() / "Downloads"
        descargas.mkdir(exist_ok=True)
        ruta = descargas / nombre
        ruta.write_bytes(contenido)
        os.startfile(str(ruta))
        snack_fn(f"✓ Guardado en Descargas: {nombre}")
    else:
        # Web (Render / Linux): descarga via data URL en el navegador
        b64 = base64.b64encode(contenido).decode()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        page.launch_url(f"data:{mime};base64,{b64}")
        snack_fn(f"✓ Descargando: {nombre}")


# ══════════════════════════════════════════════
#  HELPERS UI
# ══════════════════════════════════════════════

def ornamento(color=ACCENT):
    return ft.Row(
        controls=[
            ft.Container(width=44, height=1, bgcolor=color),
            ft.Container(width=8, height=8, bgcolor=color,
                         border_radius=4, margin=ft.margin.symmetric(horizontal=6)),
            ft.Container(width=44, height=1, bgcolor=color),
        ],
        alignment=ft.MainAxisAlignment.CENTER, spacing=0,
    )

def campo(label, form, key, on_change, hint=None, multiline=False, required=False, expand=True):
    tf = ft.TextField(
        value=form.get(key, ""),
        hint_text=hint,
        multiline=multiline,
        min_lines=2 if multiline else 1,
        max_lines=4 if multiline else 1,
        border=ft.InputBorder.UNDERLINE,
        border_color=RULE,
        focused_border_color=ACCENT,
        focused_border_width=2,
        cursor_color=ACCENT,
        text_style=ft.TextStyle(font_family=FONT_BODY, size=15, color=INK, italic=True),
        hint_style=ft.TextStyle(font_family=FONT_BODY, italic=True, color="#999"),
        content_padding=ft.padding.symmetric(vertical=8, horizontal=2),
        on_change=lambda e, k=key: (form.__setitem__(k, e.control.value), on_change()),
    )
    req_star = ft.Text(" *", color=ACCENT, size=10) if required else ft.Text("")
    return ft.Column(
        controls=[
            ft.Row(controls=[
                ft.Text(label.upper(), size=10, weight=ft.FontWeight.W_600,
                        color=INK_SOFT, font_family=FONT_BODY),
                req_star,
            ], spacing=0),
            tf,
        ],
        spacing=4, expand=expand,
    )

def tarjeta(romano, titulo, subtitulo, contenido):
    return ft.Container(
        content=ft.Column(controls=[
            ft.Row(controls=[
                ft.Container(
                    content=ft.Text(romano, size=18, italic=True,
                                    color=ACCENT, font_family=FONT_DISPLAY),
                    width=38, height=38,
                    border=ft.border.all(1, ACCENT),
                    alignment=ft.alignment.center,
                ),
                ft.Column(controls=[
                    ft.Text(titulo, size=20, weight=ft.FontWeight.W_500,
                            color=INK, font_family=FONT_DISPLAY),
                    ft.Text(subtitulo, size=11, italic=True,
                            color=INK_SOFT, font_family=FONT_BODY),
                ], spacing=2, tight=True),
            ], spacing=14),
            ft.Container(height=1, bgcolor=RULE,
                         margin=ft.margin.symmetric(vertical=12), opacity=0.6),
            contenido,
        ], spacing=0),
        bgcolor=PAPER, border=ft.border.all(1, PAPER_EDGE), padding=ft.padding.all(22),
    )


# ══════════════════════════════════════════════
#  APP PRINCIPAL
# ══════════════════════════════════════════════
def main(page: ft.Page):
    page.title   = "Solicitudes Legales"
    page.bgcolor = SURFACE
    page.padding = 0
    page.scroll  = ft.ScrollMode.AUTO

    page.fonts = {
        FONT_DISPLAY: "https://fonts.gstatic.com/s/cormorantgaramond/v16/co3YmX5slCNuHLi8bLeY9MK7whWMhyjornGV.ttf",
        FONT_BODY:    "https://fonts.gstatic.com/s/sourceserifpro/v15/neIQzD-0qpwxpaWvjeD0X88SAOeauXE-pg.ttf",
    }
    page.theme = ft.Theme(font_family=FONT_BODY)

    hoy = date.today()

    # ── ESTADO DE PLANTILLA ──────────────────────
    plantilla = {"actual": "informe"}   # "informe" | "archivo" | "reporte"

    # ── FORM PLANTILLA 1: INFORME TÉCNICO ────────
    form_informe = {
        "fecha_carta":        f"{hoy.day:02d}/{hoy.month:02d}/{hoy.year}",
        "destinatario":       "Arq. Gilberto Grullón.",
        "cargo_destinatario": "Director técnico.",
        "atencion_nombre":    "Agrim. Julio Yens Seijas,",
        "atencion_cargo":     "Encargado del Depto. de Catastro",
        "asunto":             "Solicitud de Investigación Parcelaria.",
        "anexos":             ["Acto No. ", "Certificado de Inscripción Catastral No. "],
        "inscripcion_no":     "",
        "propietario":        "",
        "fecha_audiencia":    "",
        "salon":              "3",
        "piso":               "Segundo Piso",
        "lugar":              "Ciudad Judicial",
        "expediente":         "",
        "cliente":            "",
        "firmante_nombre":    "YOLANDA DE LA CRUZ VARGAS.",
        "firmante_cargo":     "Directora Legal.",
        "iniciales":          "JDCV/rr",
    }

    # ── FORM PLANTILLA 2: COPIA CERTIFICADA ──────
    form_archivo = {
        "fecha_carta":        f"{hoy.day:02d}/{hoy.month:02d}/{hoy.year}",
        "destinatario":       "Sr. Leibi Rafael Méndez Beltre.",
        "cargo_destinatario": "Encargado interino de la División de Archivos de Expedientes Legales.",
        "asunto":             "Solicitud de Copia Certificada de Expediente.",
        "nombre_cliente":     "",
        "cedula":             "",
        "descripcion_inmueble": "",
        "tipo_proceso":       "deslinde",
        "tribunal":           "",
        "firmante_nombre":    "YOLANDA DE LA CRUZ VARGAS.",
        "firmante_cargo":     "Directora Legal.",
        "iniciales":          "YDCV/rr",
    }

    # ── FORM PLANTILLA 3: REPORTE DE AUDIENCIA ───
    form_reporte = {
        "tribunal":         "",
        "expediente":       "",
        "demandantes":      "",
        "demandados":       [""],
        "asunto":           "Solicitud de Deslinde.",
        "parcela":          "",
        "caja_no":          "",
        "rol":              "",
        "fecha_audiencia":  f"{hoy.day:02d}/{hoy.month:02d}/{hoy.year}",
        "fallo":            "",
        "fecha_proxima":    "",
        "abogados":         "",
    }

    # ── FORM PLANTILLA 4: SOLICITUD DE COBROS ────
    form_cobros = {
        "fecha_carta":        f"{hoy.day:02d}/{hoy.month:02d}/{hoy.year}",
        "destinatario":       "Licda. Denisse Amadés",
        "cargo_destinatario": "Encargada dpto. de cobros.",
        "asunto":             "SI EXISTE CUENTA O PAGO.",
        "nombre_cliente":     "",
        "tipo_proceso":       "Saneamiento",
        "parcela":            "Designación Posesional No.",
        "municipio":          "",
        "provincia":          "",
        "expediente":         "",
        "cedula":             "",
        "tribunal":           "",
        "sentencia":          "",
        "fecha_sentencia":    "",
        "tribunal_original":  "",
        "firmante_nombre":    "LICDA. YOLANDA DE LA CRUZ VARGAS.",
        "firmante_cargo":     "Directora Legal.",
        "iniciales":          "YDLCV/rr",
    }

    # ── FORM PLANTILLA 5: SOLICITUD DE INSPECCIÓN ─
    form_inspeccion = {
        "fecha_carta":        f"{hoy.day:02d}/{hoy.month:02d}/{hoy.year}",
        "destinatario":       "",
        "cargo_destinatario": "Director Técnico",
        "atencion_nombre":    "",
        "atencion_cargo":     "Encargada dpto. de Inspección.",
        "asunto":             "Solicitud de Inspección",
        "cuerpo_1":           "Muy cortésmente, por medio de la presente, solicitamos que nos informe ",
        "cuerpo_2":           "",
        "firmante_nombre":    "YOLANDA DE LA CRUZ VARGAS.",
        "firmante_cargo":     "Directora Legal.",
        "iniciales":          "YDLCV/rr",
        "cc":                 "",
    }

    # ── CONTENEDOR PRINCIPAL DEL FORMULARIO ──────
    form_area = ft.Column(spacing=20, scroll=ft.ScrollMode.AUTO, expand=True)

    def snack(msg, ok=True):
        page.open(ft.SnackBar(
            content=ft.Text(msg, color=PAPER, font_family=FONT_BODY),
            bgcolor="#2E7D32" if ok else "#B71C1C",
        ))

    # ── SECCIONES PLANTILLA 1 ────────────────────
    def build_informe():
        form = form_informe

        def refrescar(): page.update()

        # Anexos dinámicos
        anexos_col = ft.Column(spacing=8)

        def build_anexos():
            anexos_col.controls.clear()
            for idx, val in enumerate(form["anexos"]):
                def make_row(i=idx, v=val):
                    tf = ft.TextField(
                        value=v,
                        hint_text="Ej: Acto No. 02/2026.",
                        border=ft.InputBorder.UNDERLINE,
                        border_color=RULE,
                        focused_border_color=ACCENT,
                        focused_border_width=2,
                        cursor_color=ACCENT,
                        text_style=ft.TextStyle(font_family=FONT_BODY, size=15, color=INK, italic=True),
                        hint_style=ft.TextStyle(font_family=FONT_BODY, italic=True, color="#999"),
                        content_padding=ft.padding.symmetric(vertical=8, horizontal=2),
                        expand=True,
                        on_change=lambda e, i_=i: (
                            form["anexos"].__setitem__(i_, e.control.value),
                        ),
                    )
                    def on_remove(e, i_=i):
                        if len(form["anexos"]) > 1:
                            form["anexos"].pop(i_)
                            build_anexos()
                            page.update()
                    return ft.Row(controls=[
                        ft.Text(f"{i+1}-", size=13, color=INK_SOFT,
                                font_family=FONT_BODY, width=28),
                        tf,
                        ft.IconButton(
                            icon=ft.Icons.REMOVE_CIRCLE_OUTLINE,
                            icon_color=ACCENT, icon_size=18,
                            tooltip="Eliminar anexo",
                            on_click=on_remove,
                        ),
                    ], spacing=6, vertical_alignment=ft.CrossAxisAlignment.CENTER)
                anexos_col.controls.append(make_row())

            def on_agregar(e):
                form["anexos"].append("")
                build_anexos()
                page.update()

            anexos_col.controls.append(
                ft.TextButton(
                    content=ft.Row(controls=[
                        ft.Icon(ft.Icons.ADD_CIRCLE_OUTLINE, color=ACCENT, size=16),
                        ft.Text("Agregar anexo", color=ACCENT, size=13, font_family=FONT_BODY),
                    ], spacing=6),
                    on_click=on_agregar,
                )
            )

        build_anexos()

        sec1 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Fecha de la carta", form, "fecha_carta", refrescar, hint="DD/MM/AAAA", required=True),
                campo("Asunto", form, "asunto", refrescar, required=True),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Destinatario", form, "destinatario", refrescar, required=True),
                campo("Cargo del destinatario", form, "cargo_destinatario", refrescar),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Nombre (Atención)", form, "atencion_nombre", refrescar),
                campo("Cargo (Atención)",  form, "atencion_cargo",  refrescar),
            ], spacing=24),
        ])

        sec2 = ft.Column(controls=[
            ft.Row(controls=[
                ft.Text("ANEXOS", size=10, weight=ft.FontWeight.W_600,
                        color=INK_SOFT, font_family=FONT_BODY),
                ft.Text(" *", color=ACCENT, size=10),
            ], spacing=0),
            anexos_col,
            ft.Container(height=14),
            campo("No. Inscripción Catastral (cuerpo)", form, "inscripcion_no", refrescar, required=True),
            ft.Container(height=14),
            campo("Nombre del propietario / propietarios", form, "propietario",
                  refrescar, multiline=True, required=True),
        ])

        sec3 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Fecha de audiencia", form, "fecha_audiencia", refrescar,
                      hint="DD/MM/AAAA", required=True),
                ft.Row(controls=[
                    campo("Salón", form, "salon", refrescar),
                    campo("Piso",  form, "piso",  refrescar),
                ], spacing=16, expand=True),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Lugar",         form, "lugar",      refrescar),
                campo("No. Expediente",form, "expediente", refrescar, required=True),
            ], spacing=24),
            ft.Container(height=14),
            campo("Nombre del cliente", form, "cliente", refrescar, multiline=True, required=True),
        ])

        sec4 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Nombre del firmante", form, "firmante_nombre", refrescar, required=True),
                campo("Cargo del firmante",  form, "firmante_cargo",  refrescar),
            ], spacing=24),
            ft.Container(height=14),
            campo("Iniciales", form, "iniciales", refrescar, expand=False),
        ])

        def on_generar(e):
            requeridos = [
                ("inscripcion_no",  "No. Inscripción Catastral"),
                ("propietario",     "Propietario"),
                ("fecha_audiencia", "Fecha de audiencia"),
                ("expediente",      "No. Expediente"),
                ("cliente",         "Cliente"),
            ]
            faltantes = [n for k, n in requeridos if not form.get(k,"").strip()]
            if not any(a.strip() for a in form.get("anexos",[])):
                faltantes.append("al menos un Anexo")
            if faltantes:
                snack("Campos requeridos: " + ", ".join(faltantes), ok=False); return
            try:
                contenido = generar_informe_tecnico(form)
                nombre = _nombre_archivo("Informe Tecnico", form.get("cliente","Solicitud"))
                _guardar_y_abrir(page, contenido, nombre, snack)
            except Exception as ex:
                snack(f"Error: {ex}", ok=False)

        boton = ft.Container(
            content=ft.Row(controls=[
                ft.Icon(ft.Icons.DOWNLOAD, color=PAPER, size=18),
                ft.Text("GENERAR DOCUMENTO WORD", size=14,
                        weight=ft.FontWeight.W_500, color=PAPER, font_family=FONT_DISPLAY),
                ft.Text("·", size=22, color=ACCENT),
            ], spacing=14, alignment=ft.MainAxisAlignment.CENTER),
            bgcolor=PANEL_INK,
            padding=ft.padding.symmetric(vertical=16, horizontal=34),
            border=ft.border.all(1, ACCENT),
            on_click=on_generar, ink=True, border_radius=2,
        )

        return [
            tarjeta("I",   "Encabezado", "Destinatario y asunto",                sec1),
            tarjeta("II",  "Inmueble",   "Inscripción catastral, anexos y propietario", sec2),
            tarjeta("III", "Audiencia",  "Fecha, lugar y expediente",             sec3),
            tarjeta("IV",  "Firma",      "Quien suscribe la solicitud",           sec4),
            ft.Container(height=6),
            ft.Container(content=boton, alignment=ft.alignment.center),
            ft.Container(
                content=ft.Text(
                    "El archivo .docx se guardará en Descargas con el formato oficial.",
                    size=11, italic=True, color=INK_SOFT, font_family=FONT_BODY,
                    text_align=ft.TextAlign.CENTER,
                ),
                alignment=ft.alignment.center,
            ),
        ]

    # ── SECCIONES PLANTILLA 2 ────────────────────
    def build_archivo():
        form = form_archivo

        def refrescar(): page.update()

        sec1 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Fecha de la carta", form, "fecha_carta", refrescar, hint="DD/MM/AAAA", required=True),
                campo("Asunto", form, "asunto", refrescar, required=True),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Destinatario (A)", form, "destinatario", refrescar, required=True),
                campo("Cargo del destinatario", form, "cargo_destinatario", refrescar),
            ], spacing=24),
        ])

        sec2 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Nombre del cliente", form, "nombre_cliente", refrescar,
                      required=True, multiline=True),
                campo("Cédula de identidad", form, "cedula", refrescar,
                      hint="000-0000000-0", required=True),
            ], spacing=24),
            ft.Container(height=14),
            campo("Descripción del inmueble",
                  form, "descripcion_inmueble", refrescar,
                  hint="Parcela No. X, del Distrito Catastral No. X…",
                  multiline=True, required=True),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Tipo de proceso", form, "tipo_proceso", refrescar,
                      hint="deslinde / desalojo / litis…", required=True),
                campo("Tribunal / Sala", form, "tribunal", refrescar,
                      multiline=True, required=True),
            ], spacing=24),
        ])

        sec3 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Nombre del firmante", form, "firmante_nombre", refrescar, required=True),
                campo("Cargo del firmante",  form, "firmante_cargo",  refrescar),
            ], spacing=24),
            ft.Container(height=14),
            campo("Iniciales", form, "iniciales", refrescar, expand=False),
        ])

        def on_generar(e):
            requeridos = [
                ("nombre_cliente",      "Nombre del cliente"),
                ("cedula",              "Cédula"),
                ("descripcion_inmueble","Descripción del inmueble"),
                ("tipo_proceso",        "Tipo de proceso"),
                ("tribunal",            "Tribunal"),
            ]
            faltantes = [n for k, n in requeridos if not form.get(k,"").strip()]
            if faltantes:
                snack("Campos requeridos: " + ", ".join(faltantes), ok=False); return
            try:
                contenido = generar_copia_certificada(form)
                nombre = _nombre_archivo("Solicitud de Archivo", form.get("nombre_cliente","Solicitud"))
                _guardar_y_abrir(page, contenido, nombre, snack)
            except Exception as ex:
                snack(f"Error: {ex}", ok=False)

        boton = ft.Container(
            content=ft.Row(controls=[
                ft.Icon(ft.Icons.DOWNLOAD, color=PAPER, size=18),
                ft.Text("GENERAR DOCUMENTO WORD", size=14,
                        weight=ft.FontWeight.W_500, color=PAPER, font_family=FONT_DISPLAY),
                ft.Text("·", size=22, color=ACCENT),
            ], spacing=14, alignment=ft.MainAxisAlignment.CENTER),
            bgcolor=PANEL_INK,
            padding=ft.padding.symmetric(vertical=16, horizontal=34),
            border=ft.border.all(1, ACCENT),
            on_click=on_generar, ink=True, border_radius=2,
        )

        return [
            tarjeta("I",   "Encabezado",  "Destinatario y asunto",         sec1),
            tarjeta("II",  "Expediente",  "Cliente, inmueble y tribunal",   sec2),
            tarjeta("III", "Firma",       "Quien suscribe la solicitud",    sec3),
            ft.Container(height=6),
            ft.Container(content=boton, alignment=ft.alignment.center),
            ft.Container(
                content=ft.Text(
                    "El archivo .docx se guardará en Descargas con el formato oficial.",
                    size=11, italic=True, color=INK_SOFT, font_family=FONT_BODY,
                    text_align=ft.TextAlign.CENTER,
                ),
                alignment=ft.alignment.center,
            ),
        ]

    # ── SECCIONES PLANTILLA 3 ────────────────────
    def build_reporte():
        form = form_reporte

        def refrescar(): page.update()

        # Demandados dinámicos
        dem_col = ft.Column(spacing=8)

        def build_demandados():
            dem_col.controls.clear()
            for idx, val in enumerate(form["demandados"]):
                def make_row(i=idx, v=val):
                    tf = ft.TextField(
                        value=v,
                        hint_text="Nombre del demandado",
                        border=ft.InputBorder.UNDERLINE,
                        border_color=RULE,
                        focused_border_color=ACCENT,
                        focused_border_width=2,
                        cursor_color=ACCENT,
                        text_style=ft.TextStyle(font_family=FONT_BODY, size=15, color=INK, italic=True),
                        hint_style=ft.TextStyle(font_family=FONT_BODY, italic=True, color="#999"),
                        content_padding=ft.padding.symmetric(vertical=8, horizontal=2),
                        expand=True,
                        on_change=lambda e, i_=i: form["demandados"].__setitem__(i_, e.control.value),
                    )
                    def on_remove(e, i_=i):
                        if len(form["demandados"]) > 1:
                            form["demandados"].pop(i_)
                            build_demandados()
                            page.update()
                    return ft.Row(controls=[
                        tf,
                        ft.IconButton(
                            icon=ft.Icons.REMOVE_CIRCLE_OUTLINE,
                            icon_color=ACCENT, icon_size=18,
                            tooltip="Eliminar",
                            on_click=on_remove,
                        ),
                    ], spacing=6, vertical_alignment=ft.CrossAxisAlignment.CENTER)
                dem_col.controls.append(make_row())

            def on_agregar(e):
                form["demandados"].append("")
                build_demandados()
                page.update()

            dem_col.controls.append(
                ft.TextButton(
                    content=ft.Row(controls=[
                        ft.Icon(ft.Icons.ADD_CIRCLE_OUTLINE, color=ACCENT, size=16),
                        ft.Text("Agregar demandado", color=ACCENT, size=13, font_family=FONT_BODY),
                    ], spacing=6),
                    on_click=on_agregar,
                )
            )

        build_demandados()

        sec1 = ft.Column(controls=[
            ft.Row(controls=[
                campo("No. Expediente", form, "expediente", refrescar, required=True),
                campo("Asunto",         form, "asunto",     refrescar, required=True),
            ], spacing=24),
            ft.Container(height=14),
            campo("Tribunal", form, "tribunal", refrescar, multiline=True, required=True),
            ft.Container(height=14),
            campo("Parcela / Designación", form, "parcela", refrescar, required=True),
        ])

        sec2 = ft.Column(controls=[
            campo("Demandante(s)", form, "demandantes", refrescar, multiline=True, required=True),
            ft.Container(height=14),
            ft.Row(controls=[
                ft.Text("DEMANDADO(S)", size=10, weight=ft.FontWeight.W_600,
                        color=INK_SOFT, font_family=FONT_BODY),
                ft.Text(" *", color=ACCENT, size=10),
            ], spacing=0),
            dem_col,
        ])

        sec3 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Fecha de la audiencia", form, "fecha_audiencia", refrescar,
                      hint="DD/MM/AAAA", required=True),
                campo("Caja No.", form, "caja_no", refrescar),
                campo("Rol",     form, "rol",     refrescar),
            ], spacing=24),
            ft.Container(height=14),
            campo("Fallo", form, "fallo", refrescar, multiline=True, required=True,
                  hint="Ej: Aplazada a los fines de que…"),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Fecha próxima audiencia", form, "fecha_proxima", refrescar,
                      hint="DD/MM/AAAA", required=True),
                campo("Abogado(s) que compareció", form, "abogados", refrescar,
                      multiline=True, required=True),
            ], spacing=24),
        ])

        def on_generar(e):
            requeridos = [
                ("tribunal",        "Tribunal"),
                ("expediente",      "No. Expediente"),
                ("demandantes",     "Demandante(s)"),
                ("parcela",         "Parcela"),
                ("fecha_audiencia", "Fecha de audiencia"),
                ("fallo",           "Fallo"),
                ("fecha_proxima",   "Fecha próxima audiencia"),
                ("abogados",        "Abogado(s)"),
            ]
            faltantes = [n for k, n in requeridos if not form.get(k,"").strip()]
            if not any(x.strip() for x in form.get("demandados",[])):
                faltantes.append("Demandado(s)")
            if faltantes:
                snack("Campos requeridos: " + ", ".join(faltantes), ok=False); return
            try:
                contenido = generar_reporte_audiencia(form)
                # Reporte: usa el primer demandado o expediente como identificador
                dem = next((x.strip() for x in form.get("demandados",[]) if x.strip()), "")
                ident = dem or form.get("expediente","Reporte")
                nombre = _nombre_archivo("Reporte de Audiencia", ident)
                _guardar_y_abrir(page, contenido, nombre, snack)
            except Exception as ex:
                snack(f"Error: {ex}", ok=False)

        boton = ft.Container(
            content=ft.Row(controls=[
                ft.Icon(ft.Icons.DOWNLOAD, color=PAPER, size=18),
                ft.Text("GENERAR DOCUMENTO WORD", size=14,
                        weight=ft.FontWeight.W_500, color=PAPER, font_family=FONT_DISPLAY),
                ft.Text("·", size=22, color=ACCENT),
            ], spacing=14, alignment=ft.MainAxisAlignment.CENTER),
            bgcolor=PANEL_INK,
            padding=ft.padding.symmetric(vertical=16, horizontal=34),
            border=ft.border.all(1, ACCENT),
            on_click=on_generar, ink=True, border_radius=2,
        )

        return [
            tarjeta("I",   "Expediente",  "Tribunal, partes y parcela",       sec1),
            tarjeta("II",  "Partes",      "Demandante y demandado(s)",         sec2),
            tarjeta("III", "Audiencia",   "Fecha, fallo y próxima audiencia",  sec3),
            ft.Container(height=6),
            ft.Container(content=boton, alignment=ft.alignment.center),
            ft.Container(
                content=ft.Text(
                    "El archivo .docx se guardará en Descargas con el formato oficial.",
                    size=11, italic=True, color=INK_SOFT, font_family=FONT_BODY,
                    text_align=ft.TextAlign.CENTER,
                ),
                alignment=ft.alignment.center,
            ),
        ]

    # ── SECCIONES PLANTILLA 4 ────────────────────
    def build_cobros():
        form = form_cobros
        def refrescar(): page.update()

        sec1 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Fecha de la carta",    form, "fecha_carta",        refrescar, hint="DD/MM/AAAA", required=True),
                campo("Asunto",               form, "asunto",             refrescar, required=True),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Destinatario (A LA)",  form, "destinatario",       refrescar, required=True),
                campo("Cargo del destinatario", form, "cargo_destinatario", refrescar),
            ], spacing=24),
        ])

        sec2 = ft.Column(controls=[
            campo("Nombre del cliente", form, "nombre_cliente", refrescar,
                  required=True, multiline=True,
                  hint="Nombre completo tal como aparece en el expediente"),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Tipo de proceso", form, "tipo_proceso", refrescar,
                      hint="Saneamiento / Deslinde…", required=True),
                campo("Parcela / Designación", form, "parcela", refrescar,
                      hint="Designación Posesional No. …", required=True),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Municipio", form, "municipio", refrescar, required=True),
                campo("Provincia", form, "provincia", refrescar, required=True),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("No. Expediente", form, "expediente", refrescar, required=True),
                campo("Cédula",         form, "cedula",     refrescar,
                      hint="000-0000000-0", required=True),
            ], spacing=24),
        ])

        sec3 = ft.Column(controls=[
            campo("Tribunal", form, "tribunal", refrescar, multiline=True, required=True),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("No. Sentencia",    form, "sentencia",         refrescar, required=True),
                campo("Fecha sentencia",  form, "fecha_sentencia",   refrescar,
                      hint="01 Julio del 2022", required=True),
            ], spacing=24),
            ft.Container(height=14),
            campo("Tribunal que dictó la sentencia", form, "tribunal_original",
                  refrescar, multiline=True, required=True),
        ])

        sec4 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Nombre del firmante", form, "firmante_nombre", refrescar, required=True),
                campo("Cargo del firmante",  form, "firmante_cargo",  refrescar),
            ], spacing=24),
            ft.Container(height=14),
            campo("Iniciales", form, "iniciales", refrescar, expand=False),
        ])

        def on_generar(e):
            requeridos = [
                ("nombre_cliente",   "Nombre del cliente"),
                ("tipo_proceso",     "Tipo de proceso"),
                ("parcela",          "Parcela"),
                ("municipio",        "Municipio"),
                ("provincia",        "Provincia"),
                ("expediente",       "No. Expediente"),
                ("cedula",           "Cédula"),
                ("tribunal",         "Tribunal"),
                ("sentencia",        "No. Sentencia"),
                ("fecha_sentencia",  "Fecha sentencia"),
                ("tribunal_original","Tribunal que dictó la sentencia"),
            ]
            faltantes = [n for k, n in requeridos if not form.get(k,"").strip()]
            if faltantes:
                snack("Campos requeridos: " + ", ".join(faltantes), ok=False); return
            try:
                contenido = generar_solicitud_cobros(form)
                nombre = _nombre_archivo("Solicitud de Cobros", form.get("nombre_cliente","Solicitud"))
                _guardar_y_abrir(page, contenido, nombre, snack)
            except Exception as ex:
                snack(f"Error: {ex}", ok=False)

        boton = ft.Container(
            content=ft.Row(controls=[
                ft.Icon(ft.Icons.DOWNLOAD, color=PAPER, size=18),
                ft.Text("GENERAR DOCUMENTO WORD", size=14,
                        weight=ft.FontWeight.W_500, color=PAPER, font_family=FONT_DISPLAY),
                ft.Text("·", size=22, color=ACCENT),
            ], spacing=14, alignment=ft.MainAxisAlignment.CENTER),
            bgcolor=PANEL_INK,
            padding=ft.padding.symmetric(vertical=16, horizontal=34),
            border=ft.border.all(1, ACCENT),
            on_click=on_generar, ink=True, border_radius=2,
        )

        return [
            tarjeta("I",   "Encabezado",   "Destinatario y asunto",                    sec1),
            tarjeta("II",  "Cliente",       "Datos del cliente e inmueble",              sec2),
            tarjeta("III", "Proceso Legal", "Tribunal, sentencia y tribunal original",   sec3),
            tarjeta("IV",  "Firma",         "Quien suscribe la solicitud",               sec4),
            ft.Container(height=6),
            ft.Container(content=boton, alignment=ft.alignment.center),
            ft.Container(
                content=ft.Text(
                    "El archivo .docx se guardará en Descargas con el formato oficial.",
                    size=11, italic=True, color=INK_SOFT, font_family=FONT_BODY,
                    text_align=ft.TextAlign.CENTER,
                ),
                alignment=ft.alignment.center,
            ),
        ]

    # ── SECCIONES PLANTILLA 5 ────────────────────
    def build_inspeccion():
        form = form_inspeccion
        def refrescar(): page.update()

        sec1 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Fecha de la carta",     form, "fecha_carta",        refrescar, hint="DD/MM/AAAA", required=True),
                campo("Asunto",                form, "asunto",             refrescar, required=True),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Destinatario (A)",      form, "destinatario",       refrescar, required=True),
                campo("Cargo del destinatario",form, "cargo_destinatario", refrescar),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Atención (nombre)",     form, "atencion_nombre",    refrescar),
                campo("Atención (cargo)",      form, "atencion_cargo",     refrescar),
            ], spacing=24),
        ])

        sec2 = ft.Column(controls=[
            campo("Párrafo 1 — Cuerpo principal", form, "cuerpo_1", refrescar,
                  multiline=True, required=True,
                  hint="Muy cortésmente, por medio de la presente, solicitamos que…"),
            ft.Container(height=14),
            campo("Párrafo 2 — Adicional (opcional)", form, "cuerpo_2", refrescar,
                  multiline=True,
                  hint="Es oportuno mencionar que…"),
        ])

        sec3 = ft.Column(controls=[
            ft.Row(controls=[
                campo("Nombre del firmante", form, "firmante_nombre", refrescar, required=True),
                campo("Cargo del firmante",  form, "firmante_cargo",  refrescar),
            ], spacing=24),
            ft.Container(height=14),
            ft.Row(controls=[
                campo("Iniciales",  form, "iniciales", refrescar, expand=False),
                campo("Cc.",        form, "cc",        refrescar,
                      hint="Dirección General.", expand=True),
            ], spacing=24),
        ])

        def on_generar(e):
            requeridos = [
                ("destinatario", "Destinatario"),
                ("asunto",       "Asunto"),
                ("cuerpo_1",     "Párrafo principal"),
            ]
            faltantes = [n for k, n in requeridos if not form.get(k,"").strip()]
            if faltantes:
                snack("Campos requeridos: " + ", ".join(faltantes), ok=False); return
            try:
                contenido = generar_solicitud_inspeccion(form)
                nombre = _nombre_archivo("Solicitud de Inspeccion", form.get("destinatario","Solicitud"))
                _guardar_y_abrir(page, contenido, nombre, snack)
            except Exception as ex:
                snack(f"Error: {ex}", ok=False)

        boton = ft.Container(
            content=ft.Row(controls=[
                ft.Icon(ft.Icons.DOWNLOAD, color=PAPER, size=18),
                ft.Text("GENERAR DOCUMENTO WORD", size=14,
                        weight=ft.FontWeight.W_500, color=PAPER, font_family=FONT_DISPLAY),
                ft.Text("·", size=22, color=ACCENT),
            ], spacing=14, alignment=ft.MainAxisAlignment.CENTER),
            bgcolor=PANEL_INK,
            padding=ft.padding.symmetric(vertical=16, horizontal=34),
            border=ft.border.all(1, ACCENT),
            on_click=on_generar, ink=True, border_radius=2,
        )

        return [
            tarjeta("I",  "Encabezado", "Destinatario, atención y asunto", sec1),
            tarjeta("II", "Contenido",  "Cuerpo de la solicitud",           sec2),
            tarjeta("III","Firma",      "Quien suscribe y copia",           sec3),
            ft.Container(height=6),
            ft.Container(content=boton, alignment=ft.alignment.center),
            ft.Container(
                content=ft.Text(
                    "El archivo .docx se guardará en Descargas con el formato oficial.",
                    size=11, italic=True, color=INK_SOFT, font_family=FONT_BODY,
                    text_align=ft.TextAlign.CENTER,
                ),
                alignment=ft.alignment.center,
            ),
        ]

    # ── SELECTOR DE PLANTILLA ────────────────────
    def tab_btn(label, subtitulo, key):
        is_active = plantilla["actual"] == key

        def on_click(e, k=key):
            plantilla["actual"] = k
            rebuild_form()
            rebuild_tabs()
            page.update()

        return ft.Container(
            content=ft.Column(controls=[
                ft.Text(label, size=15, weight=ft.FontWeight.W_500,
                        color=PAPER if is_active else RULE,
                        font_family=FONT_DISPLAY),
                ft.Text(subtitulo, size=10, italic=True,
                        color=RULE if is_active else INK_SOFT,
                        font_family=FONT_BODY),
            ], spacing=2, tight=True),
            bgcolor=INK_SOFT if is_active else "transparent",
            padding=ft.padding.symmetric(vertical=12, horizontal=20),
            border=ft.border.only(bottom=ft.BorderSide(3, ACCENT if is_active else "transparent")),
            on_click=on_click, ink=True,
            border_radius=ft.border_radius.only(top_left=4, top_right=4),
        )

    KEYS = ["informe", "archivo", "reporte", "cobros", "inspeccion"]

    tabs_row = ft.Row(controls=[], spacing=4, scroll=ft.ScrollMode.AUTO)

    def on_arrow_click(e):
        idx = KEYS.index(plantilla["actual"])
        siguiente = KEYS[idx + 1] if idx < len(KEYS) - 1 else KEYS[0]
        plantilla["actual"] = siguiente
        rebuild_form()
        rebuild_tabs()
        page.update()

    scroll_hint = ft.Container(
        content=ft.Icon(ft.Icons.CHEVRON_RIGHT, color=ACCENT, size=26),
        padding=ft.padding.symmetric(horizontal=6, vertical=10),
        tooltip="Siguiente plantilla",
        on_click=on_arrow_click,
        ink=True,
        border_radius=4,
    )

    def rebuild_tabs():
        tabs_row.controls = [
            tab_btn("Informe Técnico",        "Catastro · Investigación Parcelaria", "informe"),
            tab_btn("Solicitud de Archivo",   "Archivo · Expediente",                "archivo"),
            tab_btn("Reporte de Audiencia",   "Audiencia · Fallo y próxima fecha",   "reporte"),
            tab_btn("Solicitud de Cobros",    "Cobros · Verificación de pago",       "cobros"),
            tab_btn("Solicitud de Inspección","Inspección · Cuerpo libre",           "inspeccion"),
        ]
        scroll_hint.visible = plantilla["actual"] != KEYS[-1]

    rebuild_tabs()

    subtitulos = {
        "informe":    ("Nueva solicitud al Catastro",     "Complete los campos para generar la solicitud de Informe Técnico."),
        "archivo":    ("Solicitud de Archivo",            "Complete los campos para generar la solicitud de Copia Certificada."),
        "reporte":    ("Reporte de Audiencia",            "Complete los campos para generar el reporte de la audiencia."),
        "cobros":     ("Solicitud de Cobros",             "Complete los campos para verificar existencia de cuenta o pago."),
        "inspeccion": ("Solicitud de Inspección",         "Redacte el cuerpo libremente y genere el documento oficial."),
    }

    header_titulo = ft.Text("", size=30, weight=ft.FontWeight.W_500, color=INK, font_family=FONT_DISPLAY)
    header_sub    = ft.Text("", size=13, italic=True, color=INK_SOFT, font_family=FONT_BODY)

    def rebuild_form():
        p = plantilla["actual"]
        header_titulo.value, header_sub.value = subtitulos[p]
        if p == "informe":
            secciones = build_informe()
        elif p == "archivo":
            secciones = build_archivo()
        elif p == "reporte":
            secciones = build_reporte()
        elif p == "cobros":
            secciones = build_cobros()
        else:
            secciones = build_inspeccion()
        form_area.controls = [
            ft.Container(
                content=ft.Column(controls=[
                    ornamento(ACCENT),
                    ft.Container(height=10),
                    ft.Text("EXPEDIENTE", size=12, color=INK_SOFT, font_family=FONT_BODY),
                    header_titulo,
                    header_sub,
                ], horizontal_alignment=ft.CrossAxisAlignment.START),
                padding=ft.padding.only(bottom=6),
            ),
            *secciones,
        ]

    rebuild_form()

    # ── TOPBAR ──────────────────────────────────
    topbar = ft.Container(
        content=ft.Column(controls=[
            ft.Row(controls=[
                ft.Column(controls=[
                    ft.Text("SISTEMA DE SOLICITUDES", size=10, italic=True,
                            color=RULE, font_family=FONT_BODY),
                    ft.Text("Automatización Legal", size=22,
                            weight=ft.FontWeight.W_500, color=PAPER,
                            font_family=FONT_DISPLAY),
                ], spacing=2, tight=True),
            ], alignment=ft.MainAxisAlignment.START),
            ft.Row(
                controls=[
                    ft.Container(content=tabs_row, expand=True),
                    scroll_hint,
                ],
                spacing=0,
                vertical_alignment=ft.CrossAxisAlignment.END,
            ),
        ], spacing=0),
        bgcolor=PANEL_INK,
        padding=ft.padding.only(top=18, bottom=0, left=36, right=8),
        border=ft.border.only(bottom=ft.BorderSide(3, ACCENT)),
    )

    layout = ft.Container(
        content=form_area,
        padding=ft.padding.symmetric(horizontal=48, vertical=24),
        expand=True,
    )

    page.add(ft.Column(controls=[topbar, layout], spacing=0, expand=True))


if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 8080))
    ft.app(target=main, view=ft.AppView.WEB_BROWSER, port=port, assets_dir="assets")
